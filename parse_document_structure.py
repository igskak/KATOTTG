#!/usr/bin/env python3
"""
Скрипт для анализа структуры документа Перелик 07052025.docx
Определяет статусы субъектов и колонки с датами
"""

import docx
from collections import defaultdict
import re
from datetime import datetime
import json

def parse_docx_structure(filename):
    """
    Парсит DOCX документ и анализирует структуру таблиц
    """
    try:
        doc = docx.Document(filename)
        print(f"Документ успешно открыт: {filename}")
        print(f"Количество параграфов: {len(doc.paragraphs)}")
        print(f"Количество таблиц: {len(doc.tables)}")
        print("-" * 50)
        
        # Анализ таблиц
        table_analysis = []
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n=== ТАБЛИЦА {table_idx + 1} ===")
            
            if len(table.rows) == 0:
                print("Таблица пустая")
                continue
                
            # Получаем заголовки (первая строка)
            headers = []
            if table.rows:
                header_row = table.rows[0]
                headers = [cell.text.strip() for cell in header_row.cells]
                print(f"Заголовки: {headers}")
            
            # Анализируем содержимое таблицы
            table_data = {
                'table_index': table_idx + 1,
                'headers': headers,
                'row_count': len(table.rows),
                'column_count': len(headers) if headers else 0,
                'sample_rows': [],
                'status_columns': [],
                'date_columns': [],
                'potential_statuses': set()
            }
            
            # Анализируем несколько строк для понимания структуры
            sample_count = min(5, len(table.rows) - 1)  # Пропускаем заголовок
            for i in range(1, min(sample_count + 1, len(table.rows))):
                row = table.rows[i]
                row_data = [cell.text.strip() for cell in row.cells]
                table_data['sample_rows'].append(row_data)
                
                # Анализируем каждую ячейку
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(headers):
                        header = headers[col_idx]
                        
                        # Проверяем на статусы
                        if any(keyword in header.lower() for keyword in ['статус', 'статус', 'статус', 'статус']):
                            table_data['status_columns'].append(col_idx)
                            table_data['potential_statuses'].add(cell_text)
                        
                        # Проверяем на даты
                        if any(keyword in header.lower() for keyword in ['дата', 'дата', 'дата', 'дата', 'дата']):
                            table_data['date_columns'].append(col_idx)
                        
                        # Проверяем содержимое на даты
                        if re.match(r'\d{1,2}[./]\d{1,2}[./]\d{2,4}', cell_text):
                            if col_idx not in table_data['date_columns']:
                                table_data['date_columns'].append(col_idx)
            
            # Анализируем все строки для поиска статусов
            for i in range(1, len(table.rows)):
                row = table.rows[i]
                for col_idx, cell_text in enumerate(row.cells):
                    if col_idx in table_data['status_columns']:
                        table_data['potential_statuses'].add(cell_text.text.strip())
            
            table_analysis.append(table_data)
            
            # Выводим результаты анализа
            print(f"Количество строк: {table_data['row_count']}")
            print(f"Количество колонок: {table_data['column_count']}")
            print(f"Колонки со статусами: {[headers[i] for i in table_data['status_columns']]}")
            print(f"Колонки с датами: {[headers[i] for i in table_data['date_columns']]}")
            print(f"Найденные статусы: {sorted(table_data['potential_statuses'])}")
            
            if table_data['sample_rows']:
                print("Примеры строк:")
                for idx, row in enumerate(table_data['sample_rows'][:3]):
                    print(f"  Строка {idx + 1}: {row}")
        
        return table_analysis
        
    except Exception as e:
        print(f"Ошибка при парсинге документа: {e}")
        return None

def analyze_status_patterns(table_analysis):
    """
    Анализирует паттерны статусов across всех таблиц
    """
    print("\n" + "="*60)
    print("ОБЩИЙ АНАЛИЗ СТАТУСОВ И ДАТ")
    print("="*60)
    
    all_statuses = set()
    all_date_columns = set()
    status_by_table = {}
    
    for table in table_analysis:
        table_num = table['table_index']
        statuses = table['potential_statuses']
        date_cols = [table['headers'][i] for i in table['date_columns']]
        
        all_statuses.update(statuses)
        all_date_columns.update(date_cols)
        status_by_table[table_num] = sorted(statuses)
        
        print(f"\nТаблица {table_num}:")
        print(f"  Статусы: {sorted(statuses)}")
        print(f"  Колонки с датами: {date_cols}")
    
    print(f"\nВСЕ УНИКАЛЬНЫЕ СТАТУСЫ: {sorted(all_statuses)}")
    print(f"ВСЕ КОЛОНКИ С ДАТАМИ: {sorted(all_date_columns)}")
    
    return {
        'all_statuses': sorted(all_statuses),
        'all_date_columns': sorted(all_date_columns),
        'status_by_table': status_by_table
    }

def save_analysis_results(table_analysis, status_analysis):
    """
    Сохраняет результаты анализа в JSON файл
    """
    # Конвертируем set в list для JSON сериализации
    for table in table_analysis:
        table['potential_statuses'] = sorted(table['potential_statuses'])
        table['status_columns'] = list(table['status_columns'])
        table['date_columns'] = list(table['date_columns'])
    
    results = {
        'document_name': 'Перелик 07052025.docx',
        'analysis_date': datetime.now().isoformat(),
        'tables': table_analysis,
        'status_analysis': status_analysis
    }
    
    with open('document_structure_analysis.json', 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    print(f"\nРезультаты анализа сохранены в document_structure_analysis.json")

def analyze_table_categories(table_analysis):
    """
    Анализирует категории таблиц на основе их заголовков
    """
    print("\n" + "="*60)
    print("АНАЛИЗ КАТЕГОРИЙ ТАБЛИЦ")
    print("="*60)
    
    categories = {}
    
    for table in table_analysis:
        headers = table['headers']
        table_num = table['table_index']
        
        # Определяем категорию на основе заголовков
        category = "Неизвестная категория"
        
        if any("можливості бойових дій" in header for header in headers):
            category = "Можливість бойових дій"
        elif any("бойових дій" in header for header in headers):
            category = "Бойові дії"
        elif any("тимчасової окупації" in header for header in headers):
            category = "Тимчасова окупація"
        
        if category not in categories:
            categories[category] = []
        categories[category].append(table_num)
        
        print(f"Таблица {table_num}: {category}")
        print(f"  Заголовки: {headers}")
        print(f"  Количество строк: {table['row_count']}")
    
    print(f"\nКАТЕГОРИИ ТАБЛИЦ:")
    for category, tables in categories.items():
        print(f"  {category}: Таблицы {tables}")
    
    return categories

if __name__ == "__main__":
    filename = "Перелик 07052025.docx"
    
    print("Начинаю анализ документа...")
    table_analysis = parse_docx_structure(filename)
    
    if table_analysis:
        status_analysis = analyze_status_patterns(table_analysis)
        categories = analyze_table_categories(table_analysis)
        save_analysis_results(table_analysis, status_analysis)
        
        # Добавляем категории в результаты
        with open('document_structure_analysis.json', 'r', encoding='utf-8') as f:
            results = json.load(f)
        
        results['categories'] = categories
        
        with open('document_structure_analysis.json', 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
    else:
        print("Не удалось проанализировать документ") 
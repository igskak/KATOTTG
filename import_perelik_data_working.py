#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Робоча версія скрипта для імпорту даних з документа Перелік 07052025 в MongoDB
Виправлено регулярне вираження для кодів територій (UA + 17 цифр)
"""

import docx
import pymongo
from pymongo import MongoClient
from urllib.parse import quote_plus
from dateutil import parser
from datetime import datetime
import json
import sys
import re
from enum import Enum

# РЯДОК ПІДКЛЮЧЕННЯ ДО MONGODB ATLAS
username = quote_plus("test")
password = quote_plus("test")
MONGO_CONNECTION_STRING = f"mongodb+srv://{username}:{password}@clusterrd4u.7180sy4.mongodb.net/?retryWrites=true&w=majority&appName=ClusterRD4U"

# Назва бази даних
DATABASE_NAME = "ua_admin_territory"

class TerritoryStatus(Enum):
    """Статуси територій згідно з документом Перелік 07052025"""
    POSSIBLE_COMBAT = "1. Території можливих бойових дій"
    ACTIVE_COMBAT = "2. Території активних бойових дій"
    ACTIVE_COMBAT_WITH_RESOURCES = "3. Території активних бойових дій, на яких функціонують державні електронні інформаційні ресурси"
    TEMPORARILY_OCCUPIED = "Тимчасово окуповані території"

def connect_to_mongodb():
    """Підключення до MongoDB Atlas"""
    try:
        print("🔌 Підключаюся до MongoDB Atlas...")
        client = MongoClient(MONGO_CONNECTION_STRING)
        client.admin.command('ping')
        print("✅ Успішно підключено до MongoDB Atlas")
        return client
    except Exception as e:
        print(f"❌ Помилка підключення до MongoDB: {e}")
        sys.exit(1)

def is_header_row(row_data):
    """
    Перевіряє, чи є рядок заголовком області/району
    """
    if not row_data or len(row_data) < 2:
        return True
    
    # Перевіряємо перші дві колонки
    first_col = row_data[0].strip() if row_data[0] else ""
    second_col = row_data[1].strip() if len(row_data) > 1 and row_data[1] else ""
    
    # Заголовки областей/районів
    header_patterns = [
        r'^\d+\.\d*\.\s*[А-ЯІЇЄ\s]+ОБЛАСТЬ$',  # 1.1. ДНІПРОПЕТРОВСЬКА ОБЛАСТЬ
        r'^\d+\.\s*[А-ЯІЇЄ\s]+ОБЛАСТЬ$',       # 1. ДНІПРОПЕТРОВСЬКА ОБЛАСТЬ
        r'^\d+\.\s*АВТОНОМНА\s+РЕСПУБЛІКА\s+КРИМ$',  # 1. АВТОНОМНА РЕСПУБЛІКА КРИМ
        r'^\d+\.\s*М\.\s*[А-ЯІЇЄ\s]+$',        # 12. М. СЕВАСТОПОЛЬ
        r'^[А-ЯІЇЄ\s]+район$',                 # Криворізький район
        r'^Найменування$',                      # Заголовок колонки
        r'^Код$',                               # Заголовок колонки
        r'^Дата\s+',                            # Заголовки дат
    ]
    
    for pattern in header_patterns:
        if re.match(pattern, first_col, re.IGNORECASE) or re.match(pattern, second_col, re.IGNORECASE):
            return True
    
    # Якщо всі колонки однакові - це заголовок
    if len(row_data) >= 2 and all(col.strip() == first_col for col in row_data[:2]):
        return True
    
    return False

def is_valid_territory_code(code):
    """
    Перевіряє, чи є код валідним кодом території
    Виправлено: UA + 17 цифр
    """
    if not code or not isinstance(code, str):
        return False
    
    # Код території має формат UA + 17 цифр
    return re.match(r'^UA\d{17}$', code.strip())

def parse_date_safely(date_str):
    """
    Безпечний парсинг дати з обробкою різних форматів
    """
    if not date_str or not isinstance(date_str, str):
        return None
    
    date_str = date_str.strip()
    if not date_str:
        return None
    
    # Видаляємо зайві символи, але зберігаємо точки та слеші
    date_str = re.sub(r'[^\d./]', '', date_str)
    
    # Якщо є кілька дат через перенос, беремо першу
    if '\n' in date_str:
        date_str = date_str.split('\n')[0]
    
    # Різні формати дат
    date_formats = [
        '%d.%m.%Y',    # 27.09.2024
        '%d/%m/%Y',    # 27/09/2024
        '%d.%m.%y',    # 27.09.24
        '%d/%m/%y',    # 27/09/24
    ]
    
    for fmt in date_formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    
    # Спробуємо парсер з dateutil
    try:
        return parser.parse(date_str, dayfirst=True)
    except:
        return None

def parse_docx_tables_improved(filename):
    """
    Покращений парсинг таблиць з DOCX документа
    """
    try:
        doc = docx.Document(filename)
        print(f"📄 Документ успішно відкрито: {filename}")
        print(f"📊 Кількість таблиць: {len(doc.tables)}")
        
        tables_data = []
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n=== ТАБЛИЦЯ {table_idx + 1} ===")
            
            if len(table.rows) == 0:
                print("Таблиця порожня")
                continue
            
            # Отримуємо заголовки
            headers = []
            if table.rows:
                header_row = table.rows[0]
                headers = [cell.text.strip() for cell in header_row.cells]
                print(f"Заголовки: {headers}")
            
            # Визначаємо статус на основі індексу таблиці
            if table_idx == 0:  # Таблица 1
                status = TerritoryStatus.POSSIBLE_COMBAT
            elif table_idx == 1:  # Таблица 2
                status = TerritoryStatus.ACTIVE_COMBAT
            elif table_idx == 2:  # Таблица 3
                status = TerritoryStatus.ACTIVE_COMBAT_WITH_RESOURCES
            elif table_idx in [3, 4]:  # Таблиці 4, 5
                status = TerritoryStatus.TEMPORARILY_OCCUPIED
            else:
                status = None
            
            print(f"Статус: {status.value if status else 'Невідомий'}")
            
            # Збираємо дані з таблиці
            table_data = {
                'table_index': table_idx + 1,
                'status': status.value if status else 'Невідомий',
                'headers': headers,
                'valid_rows': []
            }
            
            # Обробляємо всі рядки крім заголовка
            for row_idx in range(1, len(table.rows)):
                row = table.rows[row_idx]
                row_data = [cell.text.strip() for cell in row.cells]
                
                # Пропускаємо заголовки областей/районів
                if is_header_row(row_data):
                    continue
                
                # Перевіряємо, чи є валідний код території
                if len(row_data) >= 1 and is_valid_territory_code(row_data[0]):
                    table_data['valid_rows'].append(row_data)
            
            tables_data.append(table_data)
            print(f"Знайдено {len(table_data['valid_rows'])} валідних рядків")
        
        return tables_data
        
    except Exception as e:
        print(f"❌ Помилка при парсингу документа: {e}")
        return None

def find_territory_in_mongodb(client, territory_name, territory_code=None):
    """
    Пошук території в MongoDB
    """
    db = client[DATABASE_NAME]
    
    collections = [
        'level1_regions',
        'level2_raions', 
        'level3_hromadas',
        'level4_settlements',
        'level_additional_city_districts'
    ]
    
    for collection_name in collections:
        collection = db[collection_name]
        
        # Спочатку шукаємо за кодом, якщо він є
        if territory_code:
            result = collection.find_one({"_id": territory_code})
            if result:
                return result, collection_name
        
        # Пошук за точним співпадінням назви
        result = collection.find_one({"name": territory_name})
        if not result:
            # Пошук за частковим співпадінням
            result = collection.find_one({"name": {"$regex": territory_name, "$options": "i"}})
        
        if result:
            return result, collection_name
    
    return None, None

def add_status_period_to_territory(client, territory_doc, collection_name, status, start_date, end_date, 
                                  territory_code=None, table_source=None):
    """
    Додавання періоду статусу до території
    """
    db = client[DATABASE_NAME]
    collection = db[collection_name]
    
    # Визначаємо поле історії на основі статусу
    if status == TerritoryStatus.TEMPORARILY_OCCUPIED.value:
        history_field = 'occupation_history'
    elif status in [TerritoryStatus.POSSIBLE_COMBAT.value, TerritoryStatus.ACTIVE_COMBAT.value, TerritoryStatus.ACTIVE_COMBAT_WITH_RESOURCES.value]:
        history_field = 'combat_history'
    else:
        history_field = 'status_history'
    
    # Створюємо запис про статус
    status_record = {
        'status': status,
        'start_date': start_date,
        'end_date': end_date,
        'source_document': 'Перелік 07052025',
        'updated_at': datetime.now()
    }
    
    if territory_code:
        status_record['territory_code'] = territory_code
    if table_source:
        status_record['table_source'] = table_source
    
    # Отримуємо поточну історію
    if history_field in territory_doc:
        history = territory_doc[history_field]
    else:
        history = []
    
    # Перевіряємо, чи не існує вже такий запис
    existing_record = None
    for record in history:
        if (record.get('status') == status and 
            record.get('start_date') == start_date and
            record.get('source_document') == 'Перелік 07052025'):
            existing_record = record
            break
    
    if existing_record:
        print(f"⚠️  Запис вже існує для {territory_doc['name']} - статус {status}")
        return False
    
    # Додаємо новий запис до історії
    history.append(status_record)
    
    # Оновлюємо документ
    update_data = {
        "$set": {
            history_field: history,
            "current_status": status,
            "last_status_update": datetime.now()
        }
    }
    
    if start_date:
        update_data["$set"]["status_start_date"] = start_date
    
    if end_date:
        update_data["$set"]["status_end_date"] = end_date
    
    collection.update_one({"_id": territory_doc["_id"]}, update_data)
    
    print(f"✅ Додано статус '{status}' для: {territory_doc['name']}")
    return True

def import_tables_data_improved(client, tables_data):
    """
    Покращений імпорт даних з таблиць в MongoDB
    """
    print("\n🚀 Починаю імпорт даних в MongoDB...")
    
    total_imported = 0
    total_errors = 0
    not_found_territories = []
    
    for table_data in tables_data:
        table_index = table_data['table_index']
        status = table_data['status']
        headers = table_data['headers']
        valid_rows = table_data['valid_rows']
        
        print(f"\n📊 Обробляю таблицю {table_index} - {status}")
        print(f"📋 Кількість валідних рядків: {len(valid_rows)}")
        
        # Знаходимо індекси колонок
        code_idx = 0  # Код завжди перший
        name_idx = 1  # Назва завжди друга
        start_date_idx = 2  # Дата початку завжди третя
        end_date_idx = 3  # Дата кінця завжди четверта
        
        imported_in_table = 0
        errors_in_table = 0
        
        for row_idx, row_data in enumerate(valid_rows):
            if len(row_data) < 4:
                continue
            
            territory_code = row_data[code_idx] if len(row_data) > code_idx else ""
            territory_name = row_data[name_idx] if len(row_data) > name_idx else ""
            
            # Парсимо дати
            start_date = None
            end_date = None
            
            try:
                if len(row_data) > start_date_idx and row_data[start_date_idx]:
                    start_date = parse_date_safely(row_data[start_date_idx])
                
                if len(row_data) > end_date_idx and row_data[end_date_idx]:
                    end_date = parse_date_safely(row_data[end_date_idx])
            except Exception as e:
                print(f"⚠️  Помилка парсингу дати для {territory_name}: {e}")
                errors_in_table += 1
                continue
            
            # Шукаємо територію в MongoDB
            territory_doc, collection_name = find_territory_in_mongodb(client, territory_name, territory_code)
            
            if territory_doc:
                try:
                    success = add_status_period_to_territory(
                        client, 
                        territory_doc, 
                        collection_name, 
                        status, 
                        start_date, 
                        end_date,
                        territory_code=territory_code,
                        table_source=table_index
                    )
                    
                    if success:
                        imported_in_table += 1
                        total_imported += 1
                    else:
                        errors_in_table += 1
                        total_errors += 1
                        
                except Exception as e:
                    print(f"❌ Помилка додавання запису для {territory_name}: {e}")
                    errors_in_table += 1
                    total_errors += 1
            else:
                not_found_territories.append({
                    'name': territory_name,
                    'code': territory_code,
                    'status': status,
                    'table': table_index
                })
                errors_in_table += 1
                total_errors += 1
        
        print(f"📊 Таблиця {table_index}: імпортовано {imported_in_table}, помилок {errors_in_table}")
    
    # Виводимо підсумки
    print(f"\n🎯 ПІДСУМКИ ІМПОРТУ:")
    print(f"✅ Успішно імпортовано: {total_imported}")
    print(f"❌ Помилок: {total_errors}")
    
    if not_found_territories:
        print(f"\n⚠️  НЕ ЗНАЙДЕНІ ТЕРИТОРІЇ ({len(not_found_territories)}):")
        for territory in not_found_territories[:10]:  # Показуємо перші 10
            print(f"  • {territory['name']} (код: {territory['code']}) - {territory['status']}")
        
        if len(not_found_territories) > 10:
            print(f"  ... та ще {len(not_found_territories) - 10} територій")
        
        # Зберігаємо список не знайдених територій
        with open('not_found_territories_working.json', 'w', encoding='utf-8') as f:
            json.dump(not_found_territories, f, ensure_ascii=False, indent=2, default=str)
        print(f"📄 Список збережено в not_found_territories_working.json")
    
    return total_imported, total_errors

def show_import_statistics(client):
    """
    Показ статистики після імпорту
    """
    db = client[DATABASE_NAME]
    
    collections = [
        'level1_regions',
        'level2_raions', 
        'level3_hromadas',
        'level4_settlements',
        'level_additional_city_districts'
    ]
    
    print("\n📈 СТАТИСТИКА ПІСЛЯ ІМПОРТУ:")
    print("-" * 40)
    
    total_with_status = 0
    status_counts = {}
    
    for collection_name in collections:
        collection = db[collection_name]
        
        # Території з історією статусів
        with_status = collection.count_documents({
            "$or": [
                {"occupation_history": {"$exists": True}},
                {"combat_history": {"$exists": True}},
                {"status_history": {"$exists": True}}
            ]
        })
        total_with_status += with_status
        
        print(f"{collection_name}: {with_status} територій зі статусами")
        
        # Підрахунок по статусах
        territories = collection.find({
            "$or": [
                {"occupation_history": {"$exists": True}},
                {"combat_history": {"$exists": True}},
                {"status_history": {"$exists": True}}
            ]
        })
        
        for territory in territories:
            for history_field in ['occupation_history', 'combat_history', 'status_history']:
                if history_field in territory:
                    for period in territory[history_field]:
                        status = period.get('status', 'Невідомий')
                        if status not in status_counts:
                            status_counts[status] = 0
                        status_counts[status] += 1
    
    print(f"\n📊 ЗАГАЛЬНА СТАТИСТИКА:")
    print(f"Територій зі статусами: {total_with_status}")
    
    print(f"\n📋 РОЗПОДІЛ ПО СТАТУСАХ:")
    for status, count in sorted(status_counts.items()):
        print(f"  {status}: {count}")

def main():
    """Головна функція"""
    print("🚀 ІМПОРТ ДАНИХ З ДОКУМЕНТА ПЕРЕЛІК 07052025 (РОБОЧА ВЕРСІЯ)")
    print("=" * 70)
    
    # Підключаємося до MongoDB
    client = connect_to_mongodb()
    
    # Парсимо документ
    filename = "Перелик 07052025.docx"
    tables_data = parse_docx_tables_improved(filename)
    
    if not tables_data:
        print("❌ Не вдалося отримати дані з документа")
        return
    
    # Підтвердження імпорту
    print(f"\n📋 Підготовлено до імпорту:")
    total_valid_rows = 0
    for table in tables_data:
        valid_count = len(table['valid_rows'])
        total_valid_rows += valid_count
        print(f"  Таблиця {table['table_index']}: {table['status']} - {valid_count} валідних рядків")
    
    print(f"📊 Загалом валідних рядків: {total_valid_rows}")
    
    confirm = input("\n🤔 Продовжити імпорт? (y/N): ").strip().lower()
    if confirm != 'y':
        print("❌ Імпорт скасовано")
        return
    
    # Імпортуємо дані
    imported, errors = import_tables_data_improved(client, tables_data)
    
    # Показуємо статистику
    show_import_statistics(client)
    
    print(f"\n✅ Імпорт завершено!")
    print(f"📊 Успішно імпортовано: {imported}")
    print(f"❌ Помилок: {errors}")
    
    # Закриваємо з'єднання
    client.close()
    print("🔌 З'єднання з MongoDB закрито")

if __name__ == "__main__":
    main() 
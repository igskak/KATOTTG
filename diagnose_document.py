#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Діагностичний скрипт для аналізу структури документа Перелік 07052025
"""

import docx
import re

def analyze_document_structure(filename):
    """
    Аналіз структури документа
    """
    try:
        doc = docx.Document(filename)
        print(f"📄 Документ: {filename}")
        print(f"📊 Кількість таблиць: {len(doc.tables)}")
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n=== ТАБЛИЦЯ {table_idx + 1} ===")
            print(f"Кількість рядків: {len(table.rows)}")
            
            if len(table.rows) == 0:
                print("Таблиця порожня")
                continue
            
            # Аналізуємо перші 10 рядків
            for row_idx in range(min(10, len(table.rows))):
                row = table.rows[row_idx]
                row_data = [cell.text.strip() for cell in row.cells]
                
                print(f"\nРядок {row_idx + 1}:")
                print(f"  Кількість комірок: {len(row_data)}")
                print(f"  Дані: {row_data}")
                
                # Перевіряємо, чи є код території
                if row_data and len(row_data) > 0:
                    first_cell = row_data[0]
                    if re.match(r'^UA\d{12}$', first_cell):
                        print(f"  ✅ Валідний код території: {first_cell}")
                    else:
                        print(f"  ❌ Не валідний код: {first_cell}")
                
                # Перевіряємо, чи є дати
                if len(row_data) >= 3:
                    date_cells = row_data[2:4]
                    for i, date_cell in enumerate(date_cells):
                        if date_cell and re.search(r'\d{2}[./]\d{2}[./]\d{4}', date_cell):
                            print(f"  📅 Дата в комірці {i+3}: {date_cell}")
                        elif date_cell:
                            print(f"  ❓ Комірка {i+3}: {date_cell}")
            
            # Показуємо статистику по кодах територій
            print(f"\n📊 СТАТИСТИКА ТАБЛИЦІ {table_idx + 1}:")
            valid_codes = 0
            total_rows = 0
            
            for row_idx in range(1, len(table.rows)):  # Пропускаємо заголовок
                row = table.rows[row_idx]
                row_data = [cell.text.strip() for cell in row.cells]
                
                if row_data and len(row_data) > 0:
                    total_rows += 1
                    first_cell = row_data[0]
                    if re.match(r'^UA\d{12}$', first_cell):
                        valid_codes += 1
            
            print(f"  Всього рядків (без заголовка): {total_rows}")
            print(f"  Валідних кодів: {valid_codes}")
            print(f"  Відсоток валідних: {valid_codes/total_rows*100:.1f}%" if total_rows > 0 else "  Відсоток валідних: 0%")
        
    except Exception as e:
        print(f"❌ Помилка при аналізі документа: {e}")

def test_regex_patterns():
    """
    Тестування регулярних виразів
    """
    print("\n🧪 ТЕСТУВАННЯ РЕГУЛЯРНИХ ВИРАЗІВ:")
    
    test_codes = [
        "UA12060090000074553",
        "UA12060130000012028", 
        "Криворізький район",
        "1. ДНІПРОПЕТРОВСЬКА ОБЛАСТЬ",
        "12. М. СЕВАСТОПОЛЬ",
        "Найменування",
        "Код"
    ]
    
    for test_code in test_codes:
        is_valid = re.match(r'^UA\d{12}$', test_code)
        print(f"  {test_code}: {'✅ Валідний' if is_valid else '❌ Не валідний'}")
    
    # Тестування заголовків
    print("\n📋 ТЕСТУВАННЯ ЗАГОЛОВКІВ:")
    header_patterns = [
        r'^\d+\.\s*[А-ЯІЇЄ\s]+ОБЛАСТЬ$',
        r'^\d+\.\s*М\.\s*[А-ЯІЇЄ\s]+$',
        r'^[А-ЯІЇЄ\s]+район$',
        r'^Найменування$',
        r'^Код$',
        r'^Дата\s+'
    ]
    
    test_headers = [
        "1. ДНІПРОПЕТРОВСЬКА ОБЛАСТЬ",
        "12. М. СЕВАСТОПОЛЬ", 
        "Криворізький район",
        "Найменування",
        "Код",
        "Дата початку"
    ]
    
    for header in test_headers:
        is_header = False
        for pattern in header_patterns:
            if re.match(pattern, header, re.IGNORECASE):
                is_header = True
                break
        print(f"  {header}: {'📋 Заголовок' if is_header else '📄 Дані'}")

if __name__ == "__main__":
    print("🔍 ДІАГНОСТИКА ДОКУМЕНТА ПЕРЕЛІК 07052025")
    print("=" * 50)
    
    filename = "Перелик 07052025.docx"
    analyze_document_structure(filename)
    test_regex_patterns() 
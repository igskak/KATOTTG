#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для парсингу даних про окупацію територій з документа Word та оновлення MongoDB
Автор: AI Assistant
Дата: 2025
"""

import pandas as pd
import pymongo
from pymongo import MongoClient
import sys
import os
from urllib.parse import quote_plus
from docx import Document
from dateutil import parser
from datetime import datetime
import re

# РЯДОК ПІДКЛЮЧЕННЯ ДО MONGODB ATLAS
username = quote_plus("test")
password = quote_plus("test")
MONGO_CONNECTION_STRING = f"mongodb+srv://{username}:{password}@clusterrd4u.7180sy4.mongodb.net/?retryWrites=true&w=majority&appName=ClusterRD4U"

# Назва бази даних
DATABASE_NAME = "ua_admin_territory"

def connect_to_mongodb():
    """Підключення до MongoDB Atlas"""
    try:
        print("🔌 Підключаюся до MongoDB Atlas...")
        client = MongoClient(MONGO_CONNECTION_STRING)
        client.admin.command('ping')
        print("✅ Успішно підключено до MongoDB Atlas")
        return client
    except Exception as e:
        print(f"❌ Помилка підключення до MongoDB: {e}")
        sys.exit(1)

def parse_word_document(file_path):
    """
    Парсинг документа Word з даними про окупацію
    Очікуємо структуру: Назва території | Дата початку окупації | Дата кінця окупації | Статус
    """
    try:
        print(f"📖 Читаю документ Word: {file_path}")
        doc = Document(file_path)
        
        occupation_data = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Спроба розпарсити рядок з даними про окупацію
            # Очікуємо формат: "Назва території | Дата початку | Дата кінця | Статус"
            parts = [part.strip() for part in text.split('|')]
            
            if len(parts) >= 3:
                territory_name = parts[0]
                start_date_str = parts[1]
                end_date_str = parts[2] if len(parts) > 2 else ""
                status = parts[3] if len(parts) > 3 else "Тимчасово окупована"
                
                # Парсинг дат
                start_date = None
                end_date = None
                
                try:
                    if start_date_str and start_date_str.lower() not in ['', 'н/д', 'не вказано']:
                        # Очищаем от переносов строк и лишних пробелов
                        clean_start_date = start_date_str.replace('\n', ' ').replace('\r', ' ').strip()
                        if clean_start_date:
                            start_date = parser.parse(clean_start_date, dayfirst=True)
                except:
                    print(f"⚠️  Не вдалося розпарсити дату початку: {start_date_str}")
                
                try:
                    if end_date_str and end_date_str.lower() not in ['', 'н/д', 'не вказано']:
                        # Очищаем от переносов строк и лишних пробелов
                        clean_end_date = end_date_str.replace('\n', ' ').replace('\r', ' ').strip()
                        if clean_end_date:
                            end_date = parser.parse(clean_end_date, dayfirst=True)
                except:
                    print(f"⚠️  Не вдалося розпарсити дату кінця: {end_date_str}")
                
                occupation_data.append({
                    'territory_name': territory_name,
                    'start_date': start_date,
                    'end_date': end_date,
                    'status': status
                })
        
        print(f"✅ Знайдено {len(occupation_data)} записів про окупацію")
        return occupation_data
        
    except Exception as e:
        print(f"❌ Помилка читання документа Word: {e}")
        sys.exit(1)

def parse_word_table(file_path):
    """
    Парсинг таблицы Word с колонками: Код | Найменування | Дата виникнення можливості бойових дій | Дата припинення можливості бойових дій
    """
    try:
        print(f"📖 Читаю таблицы документа Word: {file_path}")
        doc = Document(file_path)
        occupation_data = []
        
        for table_idx, table in enumerate(doc.tables):
            print(f"  Обрабатываю таблицу {table_idx + 1} ({len(table.rows)} строк)")
            
            # Определяем индексы нужных колонок по первой строке
            header = [cell.text.strip() for cell in table.rows[0].cells]
            code_idx = None
            name_idx = None
            start_idx = None
            end_idx = None
            
            for i, col in enumerate(header):
                if 'Код' in col:
                    code_idx = i
                elif 'Найменування' in col:
                    name_idx = i
                elif 'Дата виникнення' in col:
                    start_idx = i
                elif 'Дата припинення' in col:
                    end_idx = i
            
            if None in (code_idx, name_idx, start_idx, end_idx):
                print(f"    ⚠️  Таблица {table_idx + 1} не содержит нужные колонки, пропускаю")
                continue
            
            # Обрабатываем строки, начиная со второй (после заголовка)
            for row_idx, row in enumerate(table.rows[1:], 1):
                cells = [cell.text.strip() for cell in row.cells]
                
                # Пропускаем строки-заголовки областей и районов
                if len(cells) > code_idx and cells[code_idx]:
                    code = cells[code_idx]
                    # Проверяем, что это валидный код КАТОТТГ
                    if is_valid_katottg_code(code):
                        name = cells[name_idx] if len(cells) > name_idx else ""
                        start_date_str = cells[start_idx] if len(cells) > start_idx else ""
                        end_date_str = cells[end_idx] if len(cells) > end_idx else ""
                        
                        # Парсинг дат
                        start_date = None
                        end_date = None
                        
                        try:
                            if start_date_str and start_date_str.lower() not in ['', 'н/д', 'не вказано']:
                                # Очищаем от переносов строк и лишних пробелов
                                clean_start_date = start_date_str.replace('\n', ' ').replace('\r', ' ').strip()
                                if clean_start_date:
                                    start_date = parser.parse(clean_start_date, dayfirst=True)
                        except:
                            print(f"    ⚠️  Не вдалося розпарсити дату початку: {start_date_str}")
                        
                        try:
                            if end_date_str and end_date_str.lower() not in ['', 'н/д', 'не вказано']:
                                # Очищаем от переносов строк и лишних пробелов
                                clean_end_date = end_date_str.replace('\n', ' ').replace('\r', ' ').strip()
                                if clean_end_date:
                                    end_date = parser.parse(clean_end_date, dayfirst=True)
                        except:
                            print(f"    ⚠️  Не вдалося розпарсити дату кінця: {end_date_str}")
                        
                        occupation_data.append({
                            'code': code,
                            'territory_name': name,
                            'start_date': start_date,
                            'end_date': end_date,
                            'status': 'Тимчасово окупована' if not end_date else 'Звільнена'
                        })
                    else:
                        # Это заголовок области или района, пропускаем
                        pass
        
        print(f"✅ Знайдено {len(occupation_data)} записів у таблицях")
        return occupation_data
        
    except Exception as e:
        print(f"❌ Помилка читання таблиці Word: {e}")
        sys.exit(1)

def find_territory_in_mongodb(db, territory_name):
    """
    Пошук території в MongoDB за назвою
    Шукаємо в усіх колекціях
    """
    collections = [
        'level1_regions',
        'level2_raions', 
        'level3_hromadas',
        'level4_settlements',
        'level_additional_city_districts'
    ]
    
    for collection_name in collections:
        collection = db[collection_name]
        
        # Пошук за точним співпадінням назви
        result = collection.find_one({"name": territory_name})
        if result:
            return result, collection_name
            
        # Пошук за частковим співпадінням (якщо точний пошук не дав результатів)
        result = collection.find_one({"name": {"$regex": territory_name, "$options": "i"}})
        if result:
            return result, collection_name
    
    return None, None

def find_territory_by_code(db, code):
    collections = [
        'level1_regions',
        'level2_raions', 
        'level3_hromadas',
        'level4_settlements',
        'level_additional_city_districts'
    ]
    for collection_name in collections:
        collection = db[collection_name]
        result = collection.find_one({"_id": code})
        if result:
            return result, collection_name
    return None, None

def update_occupation_history(db, territory_doc, collection_name, occupation_data):
    """
    Оновлення історії окупації для території
    Підтримує множинні періоди окупації
    """
    collection = db[collection_name]
    
    # Створюємо запис про окупацію
    occupation_record = {
        'start_date': occupation_data['start_date'],
        'end_date': occupation_data['end_date'],
        'status': occupation_data['status'],
        'updated_at': datetime.now()
    }
    
    # Отримуємо поточну історію окупації
    current_doc = collection.find_one({"_id": territory_doc["_id"]})
    
    if current_doc and 'occupation_history' in current_doc:
        occupation_history = current_doc['occupation_history']
    else:
        occupation_history = []
    
    # Додаємо новий запис до історії
    occupation_history.append(occupation_record)
    
    # Оновлюємо документ
    update_data = {
        "$set": {
            "occupation_history": occupation_history,
            "current_occupation_status": occupation_data['status'],
            "last_occupation_update": datetime.now()
        }
    }
    
    # Якщо є дата початку окупації, додаємо її
    if occupation_data['start_date']:
        update_data["$set"]["occupation_start_date"] = occupation_data['start_date']
    
    # Якщо є дата кінця окупації, додаємо її
    if occupation_data['end_date']:
        update_data["$set"]["occupation_end_date"] = occupation_data['end_date']
    
    collection.update_one({"_id": territory_doc["_id"]}, update_data)
    
    print(f"✅ Оновлено історію окупації для: {territory_doc['name']}")

def update_occupation_history_by_code(db, territory_doc, collection_name, occupation_data):
    collection = db[collection_name]
    occupation_record = {
        'start_date': occupation_data['start_date'],
        'end_date': occupation_data['end_date'],
        'status': occupation_data['status'],
        'updated_at': datetime.now()
    }
    current_doc = collection.find_one({"_id": territory_doc["_id"]})
    if current_doc and 'occupation_history' in current_doc:
        occupation_history = current_doc['occupation_history']
    else:
        occupation_history = []
    occupation_history.append(occupation_record)
    update_data = {
        "$set": {
            "occupation_history": occupation_history,
            "current_occupation_status": occupation_data['status'],
            "last_occupation_update": datetime.now(),
            "occupation_start_date": occupation_data['start_date'],
            "occupation_end_date": occupation_data['end_date']
        }
    }
    collection.update_one({"_id": territory_doc["_id"]}, update_data)
    print(f"✅ Оновлено історію окупації для: {territory_doc['name']} ({territory_doc['_id']})")

def process_occupation_data(client, occupation_data):
    """Обробка даних про окупацію та оновлення MongoDB"""
    db = client[DATABASE_NAME]
    
    updated_count = 0
    not_found_count = 0
    
    print("🚀 Починаю обробку даних про окупацію...")
    
    for data in occupation_data:
        territory_name = data['territory_name']
        
        # Шукаємо територію в MongoDB
        territory_doc, collection_name = find_territory_in_mongodb(db, territory_name)
        
        if territory_doc:
            update_occupation_history(db, territory_doc, collection_name, data)
            updated_count += 1
        else:
            print(f"⚠️  Територію не знайдено: {territory_name}")
            not_found_count += 1
    
    print(f"\n📈 Статистика обробки:")
    print(f"   Оновлено: {updated_count} територій")
    print(f"   Не знайдено: {not_found_count} територій")

def process_occupation_data_by_code(client, occupation_data):
    db = client[DATABASE_NAME]
    updated_count = 0
    not_found_count = 0
    print("🚀 Починаю обробку даних про окупацію (по коду)...")
    for data in occupation_data:
        code = data['code']
        territory_doc, collection_name = find_territory_by_code(db, code)
        if territory_doc:
            update_occupation_history_by_code(db, territory_doc, collection_name, data)
            updated_count += 1
        else:
            print(f"⚠️  Територію не знайдено за кодом: {code} ({data['territory_name']})")
            not_found_count += 1
    print(f"\n📈 Статистика обробки:")
    print(f"   Оновлено: {updated_count} територій")
    print(f"   Не знайдено: {not_found_count} територій")

def query_occupation_status(client, query_date):
    """
    Запит статусу окупації на конкретну дату
    """
    db = client[DATABASE_NAME]
    
    collections = [
        'level1_regions',
        'level2_raions', 
        'level3_hromadas',
        'level4_settlements',
        'level_additional_city_districts'
    ]
    
    occupied_territories = []
    
    for collection_name in collections:
        collection = db[collection_name]
        
        # Шукаємо території, які були окуповані на вказану дату
        pipeline = [
            {
                "$match": {
                    "occupation_history": {"$exists": True}
                }
            },
            {
                "$addFields": {
                    "occupation_status_on_date": {
                        "$filter": {
                            "input": "$occupation_history",
                            "as": "period",
                            "cond": {
                                "$and": [
                                    {
                                        "$lte": [
                                            "$$period.start_date",
                                            query_date
                                        ]
                                    },
                                    {
                                        "$or": [
                                            {"$eq": ["$$period.end_date", None]},
                                            {"$gte": ["$$period.end_date", query_date]}
                                        ]
                                    }
                                ]
                            }
                        }
                    }
                }
            },
            {
                "$match": {
                    "occupation_status_on_date": {"$ne": []}
                }
            }
        ]
        
        results = list(collection.aggregate(pipeline))
        
        for result in results:
            occupied_territories.append({
                'name': result['name'],
                'category': result['category'],
                'collection': collection_name,
                'occupation_periods': result['occupation_status_on_date']
            })
    
    return occupied_territories

def is_valid_katottg_code(code):
    """Проверяет, что код КАТОТТГ валидный: начинается с UA и далее только цифры, длина 18-20 символов"""
    return bool(re.fullmatch(r'UA\d{17,19}', code))

def main():
    """Головна функція"""
    print("=" * 60)
    print("🚀 СКРИПТ ПАРСИНГУ ДАНИХ ПРО ОКУПАЦІЮ")
    print("=" * 60)
    
    # Перевіряємо аргументи командного рядка
    if len(sys.argv) < 2:
        print("❌ Використання: python3 parse_occupation_data.py <шлях_до_документа.docx>")
        print("   Або: python3 parse_occupation_data.py --query <дата>")
        sys.exit(1)
    
    # Підключаємося до MongoDB
    client = connect_to_mongodb()
    
    if sys.argv[1] == "--query":
        if len(sys.argv) < 3:
            print("❌ Вкажіть дату для запиту: python3 parse_occupation_data.py --query <дата>")
            sys.exit(1)
        
        try:
            query_date = parser.parse(sys.argv[2], dayfirst=True)
            print(f"🔍 Шукаю території, окуповані на {query_date.strftime('%d.%m.%Y')}...")
            
            occupied_territories = query_occupation_status(client, query_date)
            
            print(f"\n📊 Знайдено {len(occupied_territories)} окупованих територій:")
            for territory in occupied_territories:
                print(f"   • {territory['name']} ({territory['category']})")
                for period in territory['occupation_periods']:
                    start_str = period['start_date'].strftime('%d.%m.%Y') if period['start_date'] else 'н/д'
                    end_str = period['end_date'].strftime('%d.%m.%Y') if period['end_date'] else 'н/д'
                    print(f"     Період: {start_str} - {end_str} ({period['status']})")
            
        except Exception as e:
            print(f"❌ Помилка парсингу дати: {e}")
            sys.exit(1)
    
    else:
        # Парсинг документа Word
        word_file_path = sys.argv[1]
        
        if not os.path.exists(word_file_path):
            print(f"❌ Файл '{word_file_path}' не знайдено")
            sys.exit(1)
        
        # Пробуем парсить таблицу
        occupation_data = parse_word_table(word_file_path)
        if occupation_data:
            print(f"\n📊 Знайдено {len(occupation_data)} записів для обробки")
            process_occupation_data_by_code(client, occupation_data)
        else:
            # fallback: старый парсер по параграфам
            occupation_data = parse_word_document(word_file_path)
            if occupation_data:
                process_occupation_data(client, occupation_data)
            else:
                print("❌ Не знайдено даних про окупацію в документі")
    
    # Закриваємо з'єднання
    client.close()
    print("\n🔌 З'єднання з MongoDB закрито")
    print("=" * 60)
    print("🎉 Робота скрипту завершена успішно!")

if __name__ == "__main__":
    main() 